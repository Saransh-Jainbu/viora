# """
# Text extraction and chunking service.

# Supports PDF, DOCX, and TXT files.
# Uses LangChain's RecursiveCharacterTextSplitter for chunking.
# """

# import io
# from PyPDF2 import PdfReader
# from docx import Document as DocxDocument
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from config import settings


# SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


# from typing import TypedDict

# class DocumentChunk(TypedDict):
#     text: str
#     metadata: dict

# def extract_text_with_metadata(file_bytes: bytes, filename: str) -> list[dict]:
#     """
#     Extract text with page metadata.
#     Returns a list of dicts: [{"text": "...", "page": 1}, ...]
#     """
#     ext = _get_extension(filename)
#     if ext == ".pdf":
#         return _extract_pdf_pages(file_bytes)
#     elif ext == ".docx":
#         return [{"text": _extract_docx(file_bytes), "page": 1}]
#     elif ext == ".txt":
#         return [{"text": file_bytes.decode("utf-8", errors="replace"), "page": 1}]
#     else:
#         raise ValueError(f"Unsupported extension {ext}")

# def split_into_chunks_with_metadata(
#     docs: list[dict],
#     chunk_size: int | None = None,
#     chunk_overlap: int | None = None,
# ) -> list[DocumentChunk]:
#     """
#     Split text while preserving page metadata.
#     Optimized for speed by processing larger blocks.
#     """
#     splitter = RecursiveCharacterTextSplitter(
#         chunk_size=chunk_size or settings.chunk_size,
#         chunk_overlap=chunk_overlap or settings.chunk_overlap,
#         length_function=len,
#         separators=["\n\n", "\n", ". ", " ", ""],
#     )
    
#     final_chunks: list[DocumentChunk] = []
    
#     # Process each page but skip the slow heading detection
#     for doc in docs:
#         text = doc["text"]
#         page = doc.get("page", 1)
#         sub_chunks = splitter.split_text(text)
        
#         for chunk_text in sub_chunks:
#             final_chunks.append({
#                 "text": chunk_text,
#                 "metadata": {"page": page} # Heading removed for speed and simplicity
#             })
            
#     return final_chunks

# def _extract_pdf_pages(file_bytes: bytes) -> list[dict]:
#     reader = PdfReader(io.BytesIO(file_bytes))
#     results = []
#     for i, page in enumerate(reader.pages):
#         text = page.extract_text()
#         if text:
#             results.append({"text": text, "page": i + 1})
#     return results

# def _extract_docx(file_bytes: bytes) -> str:
#     doc = DocxDocument(io.BytesIO(file_bytes))
#     paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
#     return "\n\n".join(paragraphs)

# def _get_extension(filename: str) -> str:
#     dot_index = filename.rfind(".")
#     return filename[dot_index:].lower() if dot_index != -1 else ""

"""
Text extraction and chunking service.

Supports PDF, DOCX, and TXT files.
Uses LangChain's RecursiveCharacterTextSplitter for chunking.
"""

import io
import re 

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from config import settings

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}

from typing import TypedDict

class DocumentChunk(TypedDict):
    text: str
    metadata: dict


# DOCX section extraction
def _extract_docx_with_sections(file_bytes: bytes) -> list[dict]:
    doc = DocxDocument(io.BytesIO(file_bytes))
    
    sections = []
    current_section = {"section": "Introduction", "text": ""}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name.lower()

        # detect headings
        if "heading" in style:
            if current_section["text"]:
                sections.append(current_section)

            current_section = {
                "section": text,
                "text": ""
            }
        else:
            current_section["text"] += text + "\n"

    if current_section["text"]:
        sections.append(current_section)

    return [
        {
            "text": sec["text"],
            "page": 1,
            "section": sec["section"]
        }
        for sec in sections
    ]


def extract_text_with_metadata(file_bytes: bytes, filename: str) -> list[dict]:
    """
    Extract text with page metadata.
    Returns a list of dicts: [{"text": "...", "page": 1}, ...]
    """
    ext = _get_extension(filename)

    if ext == ".pdf":
        return _extract_pdf_pages(file_bytes)

    elif ext == ".docx":
        # use section-aware extraction
        return _extract_docx_with_sections(file_bytes)

    elif ext == ".txt":
        return [{"text": file_bytes.decode("utf-8", errors="replace"), "page": 1}]

    else:
        raise ValueError(f"Unsupported extension {ext}")


def split_into_chunks_with_metadata(
    docs: list[dict],
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[DocumentChunk]:
    """
    Split text while preserving page metadata.
    Optimized for speed by processing larger blocks.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size or settings.chunk_size,
        chunk_overlap=chunk_overlap or settings.chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""],  
    )
    
    final_chunks: list[DocumentChunk] = []
    
    for doc in docs:
        text = doc["text"]
        page = doc.get("page", 1)
        section = doc.get("section", "unknown") 

        # text cleaning
        text = re.sub(r"-\n", "", text)
        text = re.sub(r"\n+", "\n", text)

        sub_chunks = splitter.split_text(text)
        
        for chunk_text in sub_chunks:
            final_chunks.append({
                "text": chunk_text,
                "metadata": {
                    "page": page,
                    "section": section  
                }
            })
            
    return final_chunks


def _extract_pdf_pages(file_bytes: bytes) -> list[dict]:
    reader = PdfReader(io.BytesIO(file_bytes))
    results = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            results.append({"text": text, "page": i + 1})
    return results

def _get_extension(filename: str) -> str:
    dot_index = filename.rfind(".")
    return filename[dot_index:].lower() if dot_index != -1 else ""